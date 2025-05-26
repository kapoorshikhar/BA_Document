from docx import Document
from docx.shared import Inches
import os

# Sample user input (simulating form submission)
company_name = "MedSource"
company_place = "Chicago, IL"
logo_path = "/mnt/data/25f69a69-1e32-4cf3-b7be-e5ceb035dedd.png"

# Create a new Word Document
doc = Document()

# Add title and content
doc.add_heading('Software Requirement Specifications', level=1)
doc.add_paragraph('For')
doc.add_heading(company_name, level=2)
doc.add_paragraph('Version 1.0')

# Add logo image if it exists
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2.0))

# Add location
doc.add_paragraph(f'Location: {company_place}')

# Save the document
output_path = "/mnt/data/SRS_Document_MedSource.docx"
doc.save(output_path)

output_path
